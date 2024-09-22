import base64
import os
import shutil
import tempfile
import time
import zipfile
from io import BytesIO

import streamlit as st
from docx import Document
from langchain_openai import ChatOpenAI
from PIL import Image

from model_langchain import HTPModel

SUPPORTED_LANGUAGES = {
    "English": "en",
    "中文": "zh"
}

LANGUAGES = {
    "en": {
        "analysis_settings": "Analysis Settings",
        "model_settings": "🍓 Model Settings",
        "batch_title": "📊 Batch Analysis",
        "language_label": "Language:",
        "select_folder": "Enter the folder path containing images:",
        "no_images_found": "No image files found in the selected folder.",
        "images_found": "{} image files found. Ready for batch analysis.",
        "start_batch_analysis": "Start Batch Analysis",
        "batch_results_summary": "Batch Analysis Results Summary",
        "download_batch_results": "Download Batch Results",
        "enter_valid_folder": "Please enter a valid folder path.",
        "error_no_api_key": "❌ Please enter your API key in the sidebar before starting the analysis.",
        "batch_instructions_title": "📋 Batch Analysis Instructions",
        "upload_images": "Upload Images for Batch Analysis",
        "images_uploaded": "{} images uploaded successfully.",
        "upload_images_prompt": "Please upload images to start batch analysis.",
    "batch_instructions": """
    **Please read the following instructions carefully before proceeding with batch analysis:**

    1. **API Key**: Ensure you have filled in your API key in the sidebar. This is crucial for the analysis to work.
    
    2. **Preparation**: 
       - Place all images you want to analyze in a single folder.
       - Make sure all images are in .jpg, .jpeg, or .png format.
    
    3. **Folder Path**: Enter the full path to the folder containing your images in the text box below.
    
    4. **Time Consideration**: Batch analysis may take a considerable amount of time, depending on the number of images. Please be patient.
    
    5. **Network and API Credits**:
       - Ensure you have a stable internet connection throughout the process.
       - Check that you have sufficient API credits for the entire batch. Each image consumes credits.
    
    6. **Results**: 
       - Use the 'Download Batch Results' button to save the full analysis results.

    **Note**: This tool is for reference only and cannot replace professional psychological evaluation. If you have concerns, please consult a qualified mental health professional.
    """,
    "welcome": "Welcome to the Batch Analysis Page",
    "batch_results": "Batch Analysis Finished, Please download the results. Successful: {} | Failed: {}",
    "download_batch_results": "Download Batch Results (ZIP)",
    "ai_disclaimer": "NOTE: AI-generated content, for reference only. Not a substitute for medical diagnosis.",
    },
    "zh": {
        "analysis_settings": "分析设置",
        "model_settings": "🍓 模型设置",
        "batch_title": "📊 批量分析",
        "language_label": "语言：",
        "select_folder": "输入包含图片的文件夹路径：",
        "no_images_found": "在选定的文件夹中未找到图片文件。",
        "images_found": "找到 {} 个图片文件，点击**开始批量分析**按钮可以开始分析。",
        "start_batch_analysis": "开始批量分析",
        "batch_results_summary": "批量分析结果摘要",
        "download_batch_results": "下载批量结果",
        "enter_valid_folder": "请输入有效的文件夹路径。",
        "error_no_api_key": "❌ 请在开始分析之前在侧边栏输入您的API密钥。",
        "batch_instructions_title": "📋 批量分析说明",
        "batch_instructions": """
        **在进行批量分析之前，请仔细阅读以下说明：**

        1. **API密钥**：确保您已在侧边栏填写了API密钥。这对分析能否进行至关重要。
        
        2. **准备工作**：
        - 将所有要分析的图片放在同一个文件夹中。
        - 确保所有图片格式为.jpg、.jpeg或.png。
        
        3. **文件夹路径**：在下方的文本框中输入包含图片的文件夹的完整路径。
        
        4. **时间考虑**：批量分析可能需要相当长的时间，具体取决于图片数量。请耐心等待。
        
        5. **网络和API额度**：
        - 确保在整个过程中网络连接稳定。
        - 检查您的API额度是否足够完成整个批次。每张图片都会消耗额度。
        
        6. **结果**：
        - 使用"下载批量结果"按钮下载完整的分析结果。

        **注意**：此工具仅供参考，不能替代专业的心理评估。如有疑虑，请咨询合格的心理健康专业人士。
        """,
        "welcome": "欢迎来到批量分析页面",
        "batch_results": "批量分析完成，请下载结果。成功: {} | 失败: {}",
        "download_batch_results": "下载批量结果 (ZIP)",
        "ai_disclaimer": "注意：本报告由AI 生成，仅供参考。不能替代医学诊断。",
        "upload_images": "上传图片进行批量分析",
        "images_uploaded": "已成功上传 {} 张图片。",
        "upload_images_prompt": "请上传图片以开始批量分析。",
    }
}

def pil_to_base64(image: Image.Image, format: str = "JPEG") -> str:
    """Convert PIL image to base64 string."""
    buffered = BytesIO()
    image.save(buffered, format=format)
    return base64.b64encode(buffered.getvalue()).decode('utf-8')

def get_text(key):
    return LANGUAGES[st.session_state['language_code']][key]

def save_results(results):
    with tempfile.TemporaryDirectory() as temp_dir:
        for result in results:
            file_name = result['file_name']
            file_name_without_ext = os.path.splitext(file_name)[0]
            
            result_folder = os.path.join(temp_dir, file_name_without_ext)
            os.makedirs(result_folder, exist_ok=True)
            
            # 保存上传的图片
            if result['image']:
                image_path = os.path.join(result_folder, file_name)
                result['image'].save(image_path)
            
            doc = Document()
            if result['success']:
                doc.add_paragraph(get_text("ai_disclaimer"))
                if result['analysis_result']['classification'] is True:
                    signal = result['analysis_result']['signal']
                    final = result['analysis_result']['final']
                    doc.add_paragraph(signal)
                    doc.add_paragraph(final)
                else:
                    signal = result['analysis_result']['fix_signal']
                    doc.add_paragraph(signal)
            else:
                doc.add_paragraph("failed")
            doc_path = os.path.join(result_folder, f"{file_name_without_ext}.docx")
            doc.save(doc_path)
        
        failed_path = os.path.join(temp_dir, "failed.txt")
        with open(failed_path, "w") as f:
            for result in results:
                if not result['success']:
                    f.write(f"{result['file_name']}\n")
                    
        zip_path = os.path.join(temp_dir, "results.zip")
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for root, _, files in os.walk(temp_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, temp_dir)
                    zipf.write(file_path, arcname)
        with open(zip_path, "rb") as f:
            zip_content = f.read()
    
    return zip_content    
        
def batch_analyze(uploaded_files):
    results = []
    
    MULTIMODAL_MODEL="gpt-4o-2024-08-06"
    TEXT_MODEL="claude-3-5-sonnet-20240620"
    
    text_model = ChatOpenAI(
        api_key=st.session_state.api_key,
        base_url=st.session_state.base_url,
        model = TEXT_MODEL,
        temperature=0.2,
        top_p = 0.75,
    )
    multimodal_model = ChatOpenAI(
        api_key=st.session_state.api_key,
        base_url=st.session_state.base_url,
        model = MULTIMODAL_MODEL,
        temperature=0.2,
        top_p = 0.75,
    )
    model = HTPModel(
        text_model=text_model,
        multimodal_model=multimodal_model,
        language=st.session_state['language_code'],
        use_cache=True
    )
    progress_bar = st.progress(0, text=f"Progressing: 0/{len(uploaded_files)}")
    start_time = time.time()
    success = 0
    for i, uploaded_file in enumerate(uploaded_files):
        try:
            image = Image.open(uploaded_file)
            image_data = pil_to_base64(image)
            
            response = model.workflow(image_path=image_data, language=st.session_state['language_code'])
            results.append({
                "file_name": uploaded_file.name,
                "analysis_result": response,
                "success": True,
                "image": image
            })
            success += 1
        except Exception as e:
            results.append({
                "file_name": uploaded_file.name,
                "analysis_result": str(e),
                "success": False,
                "image": image
            })
        
        elapsed_time = time.time() - start_time
        progress = (i + 1) / len(uploaded_files)
        estimated_total_time = elapsed_time / progress if progress > 0 else 0
        remaining_time = estimated_total_time - elapsed_time
        
        elapsed_str = time.strftime("%H:%M:%S", time.gmtime(elapsed_time))
        remaining_str = time.strftime("%H:%M:%S", time.gmtime(remaining_time))
        
        progress_bar.progress(progress, text=f"Progressing: {i + 1}/{len(uploaded_files)} | Elapsed: {elapsed_str} | Remaining: {remaining_str}")
    
    st.success(get_text("batch_results").format(success, len(uploaded_files) - success))
    
    return results, success

def sidebar() -> None: 
    """Render sidebar components."""
    st.sidebar.image("assets/logo2.png", use_column_width=True)
    
    # Analysis Settings
    st.sidebar.markdown(f"## {get_text('analysis_settings')}")
    language = st.sidebar.selectbox(
        get_text("language_label"),
        options=list(SUPPORTED_LANGUAGES.keys()),
        index=list(SUPPORTED_LANGUAGES.keys()).index(st.session_state['language']),
    )
    if language != st.session_state['language']:
        st.session_state['language'] = language
        st.session_state['language_code'] = SUPPORTED_LANGUAGES[language]
        st.rerun()
    
    uploaded_files = st.sidebar.file_uploader(get_text("upload_images"), accept_multiple_files=True, type=['png', 'jpg', 'jpeg'])
    if uploaded_files:
        st.success(get_text("images_uploaded").format(len(uploaded_files)))
    else:
        st.warning(get_text("enter_valid_folder"))
    
    # Model Settings
    st.sidebar.markdown(f"## {get_text('model_settings')}")
    base_url = st.sidebar.text_input("API Base URL", help="Base URL of the API server")
    api_key = st.sidebar.text_input("API Key", help="API Key for authentication")
    st.session_state.api_key = api_key
    st.session_state.base_url = base_url
    
    # Buttons
    st.sidebar.markdown("---")
    if st.sidebar.button(get_text("start_batch_analysis"), type="primary"):
        if not st.session_state.api_key:
            st.error(get_text("error_no_api_key"))
        else:
            results, success = batch_analyze(uploaded_files=uploaded_files)
            
            zip_content = save_results(results)
            st.download_button(
                label = get_text("download_batch_results"),
                data=zip_content,
                file_name="batch_analysis_results.zip",
                mime="application/zip"
            )
    
def batch_page():
    st.title(get_text("batch_title"))
    
    st.write(get_text("welcome"))
    
    with st.expander(get_text("batch_instructions_title"), expanded=True):
        st.markdown(get_text("batch_instructions"))
    
    
def main():
    st.set_page_config(page_title="PsychePal: Batch Analysis", page_icon="📊", layout="wide")
    
    if 'language' not in st.session_state:
        st.session_state['language'] = "中文"
        st.session_state['language_code'] = SUPPORTED_LANGUAGES[st.session_state['language']]
    
    batch_page()
    sidebar()
    

if __name__ == "__main__":
    main()